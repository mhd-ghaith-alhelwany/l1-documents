from django.core.exceptions import ValidationError
from docx import Document
from django.utils.translation import gettext_lazy as _
from django.core.files.storage import FileSystemStorage
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.shared import CT_DecimalNumber
from ast import literal_eval
from json2xml import json2xml
from json2xml.utils import readfromstring
import json
import os
import re

from docs.models import Table


class DocumentTableValidator:
    def __init__(self, file):
        self.file_name = file
        self.path = self.save_temp_file()[1:]
        self.errors = []
        self.document = None

    def validate(self):
        self.document = Document(self.path)

        if self.contains_tables():
            self.errors.append(
                ValidationError(
                    _('%(value)s does not have any tables'),
                    params={"value": self.file_name}
                )
            )
        for invalid_title in self.get_invalid_table_titles():
            self.errors.append(
                ValidationError(
                    _('\"%(value)s\" is not a valid table title'),
                    params={"value": invalid_title}
                )
            )

        for duplicated_title in self.get_duplicated_table_titles():
            self.errors.append(
                ValidationError(
                    _('\"%(value)s\" is not a unique title'),
                    params={"value": duplicated_title}
                )
            )

        for table_with_merged_cells in self.get_tables_with_merged_cells():
            self.errors.append(
                ValidationError(
                    _('table number %(value)s contains merged cells'),
                    params={"value": table_with_merged_cells}
                )
            )

        return {
            "is_valid": len(self.errors) == 0,
            "errors": self.errors
        }

    def contains_tables(self):
        return len(self.document.tables) == 0

    def get_invalid_table_titles(self):
        invalid_titles = []
        for title in get_all_table_titles(self.document):
            if not is_valid_title(title):
                invalid_titles.append(title)
        return invalid_titles

    def get_duplicated_table_titles(self):
        titles = get_all_table_titles(self.document)
        return set([x for x in titles if titles.count(x) > 1])

    def save_temp_file(self):
        fs = FileSystemStorage()
        filename = fs.save("storage/temp/document." + self.file_name.name.split(".")[-1], self.file_name)
        return fs.url(filename)

    def get_tables_with_merged_cells(self):
        tables_with_merged_cells = []
        for i in range(len(self.document.tables)):
            for row in self.document.tables[i].rows:
                for cell in row.cells:
                    cell_properties = cell._element[0].getchildren()
                    for _property in cell_properties:
                        if isinstance(_property, CT_DecimalNumber):
                            tables_with_merged_cells.append(i)
        return set(tables_with_merged_cells)


def get_all_table_titles(document):
    all_titles = []
    elements = document._element.body.getchildren()
    preceding_text = ''
    for i in range(len(elements)):
        if isinstance(elements[i], CT_P):
            paragraph_text = get_inner_paragraph_text(elements[i])
            if paragraph_text:
                preceding_text = paragraph_text
        elif isinstance(elements[i], CT_Tbl):
            all_titles.append(preceding_text)
    return all_titles


def get_inner_paragraph_text(element):
    if element.text:
        return element.text
    else:
        for child in element.getchildren():
            child_inner_text = get_inner_paragraph_text(child)
            if child_inner_text:
                return child_inner_text


def is_valid_title(title):
    return re.match(r"Table[1-9]+$", title)


class DocumentConverter:
    def __init__(self, title, tables):
        self.title = title
        self.tables = tables

    @staticmethod
    def from_file(file):
        document = Document(file)
        tables = document.tables
        titles = get_all_table_titles(document)
        tables_array = []
        for i in range(len(tables)):
            title = titles[i]
            rows = []
            for row in tables[i].rows:
                cells = []
                for cell in row.cells:
                    cells.append(cell.text)
                rows.append(cells)
            tables_array.append({
                "title": title,
                "rows": rows
            })
        return DocumentConverter(file.name, tables_array)

    @staticmethod
    def table_to_json(table):
        rows = []
        for row in literal_eval(table.json):
            cells = []
            for cell in row:
                cells.append(cell)
            rows.append({
                "cells": cells
            })
        return{
            "title": table.title,
            "rows": rows
        }

    @staticmethod
    def table_to_xml(table):
        table_json = DocumentConverter.table_to_json(table)
        table_json = json.dumps(table_json)
        data = readfromstring(table_json)
        return json2xml.Json2xml(data).to_xml()

    def get_document_file(self):
        document = Document()
        document.title = self.title
        path = "storage/temp/download.docx"
        document.save(path)
        return path

    def to_document_file(self):
        document = Document()
        for table in self.tables:
            table_array = table['rows']
            table_height = len(table_array)
            table_width = len(table_array[0])
            document.add_paragraph(table['title'])
            doc_table = document.add_table(rows=table_height, cols=table_width)
            for i in range(table_height):
                cells = doc_table.rows[i].cells
                for j in range(table_width):
                    cells[j].text = table_array[i][j]
        path = "storage/temp/" + self.title
        document.save(path)
        return os.path.abspath(path)

    @staticmethod
    def from_document_model(document_model):
        tables = Table.objects.filter(document_id=document_model).all()
        tables_json = [
            {
                "title": table.title,
                "rows": literal_eval(table.json)
            }
            for table in tables
        ]
        return DocumentConverter(document_model.file_name, tables_json)

